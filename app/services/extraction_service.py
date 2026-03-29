"""
services/extraction_service.py

Extracts raw text from uploaded CV files (PDF or DOCX).
Called after file_validators.py confirms the file is safe.

Returns a consistent (text, error) tuple matching the
(user, error) pattern used across the codebase.

Dependencies:
    pip install pymupdf python-docx
"""

import os
import re
import fitz                    # PyMuPDF — handles PDF extraction
from docx import Document      # python-docx — handles DOCX extraction


# ── Constants ─────────────────────────────────────────────────────────────────

# If extracted text is shorter than this, we treat it as a failed
# extraction — likely a scanned image PDF with no text layer.
MIN_MEANINGFUL_TEXT_LENGTH = 50


# ── Public API ────────────────────────────────────────────────────────────────

def extract_text_from_file(file_path: str) -> tuple:
    """
    Extract raw text from a PDF or DOCX file at the given path.

    Detects file type from the extension, delegates to the
    appropriate extractor, then cleans and validates the result.

    Args:
        file_path: absolute or relative path to the saved CV file

    Returns:
        (text, None)   — extraction succeeded, text is a clean string
        (None, error)  — extraction failed, error explains why
    """
    # Extension check runs FIRST — before touching the filesystem.
    # This way /fake/path/resume.txt returns "Unsupported" not "not found".
    ext = _get_extension(file_path)

    if ext not in ("pdf", "docx"):
        return None, (
            f"Unsupported file type '.{ext}'. "
            "Please upload a PDF or DOCX file."
        )

    # File existence check runs AFTER extension check
    if not os.path.exists(file_path):
        return None, "CV file not found on the server. Please upload again."

    if ext == "pdf":
        raw_text, error = _extract_from_pdf(file_path)
    else:
        raw_text, error = _extract_from_docx(file_path)

    if error:
        return None, error

    # Clean the extracted text
    cleaned_text = _clean_text(raw_text)

    # Validate the cleaned text has meaningful content
    is_valid, error = _validate_text(cleaned_text, ext)
    if not is_valid:
        return None, error

    return cleaned_text, None


# ── Private extractors ────────────────────────────────────────────────────────

def _extract_from_pdf(file_path: str) -> tuple:
    """
    Extract text from a PDF file using PyMuPDF (fitz).

    PyMuPDF reads the actual text layer of the PDF — it does NOT
    perform OCR. If the PDF is a scanned image, it will return
    empty or near-empty text, which _validate_text catches.

    Strategy:
        - Open the PDF with fitz
        - Iterate every page
        - Extract text with 'text' mode (preserves reading order)
        - Join pages with double newline to preserve structure

    Args:
        file_path: path to the PDF file

    Returns:
        (raw_text, None) or (None, error_string)
    """
    try:
        pdf_document = fitz.open(file_path)
    except Exception:
        return None, (
            "Could not open the PDF file. "
            "It may be corrupted or password-protected."
        )

    # Check if the PDF is password protected
    if pdf_document.needs_pass:
        pdf_document.close()
        return None, (
            "The PDF is password-protected. "
            "Please remove the password and upload again."
        )

    pages_text = []

    try:
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            # 'text' mode preserves reading order better than raw extraction
            page_text = page.get_text("text")
            if page_text.strip():
                pages_text.append(page_text)
    except Exception:
        return None, "Failed to read the PDF content. The file may be corrupted."
    finally:
        pdf_document.close()

    raw_text = "\n\n".join(pages_text)
    return raw_text, None


def _extract_from_docx(file_path: str) -> tuple:
    """
    Extract text from a DOCX file using python-docx.

    Reads two sources of text:
        1. Paragraphs — the main body text, headings, bullet points
        2. Table cells — some CV templates put content inside tables

    Both are joined together so no content is missed even in
    table-heavy resume templates.

    Args:
        file_path: path to the DOCX file

    Returns:
        (raw_text, None) or (None, error_string)
    """
    try:
        document = Document(file_path)
    except Exception:
        return None, (
            "Could not open the DOCX file. "
            "It may be corrupted or in an unsupported format."
        )

    text_parts = []

    # 1. Extract paragraph text (covers most CV content)
    try:
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
    except Exception:
        return None, "Failed to read DOCX paragraphs. The file may be corrupted."

    # 2. Extract table cell text (catches table-based CV layouts)
    try:
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in text_parts:
                        text_parts.append(cell_text)
    except Exception:
        # Tables are optional — log but don't fail the whole extraction
        pass

    raw_text = "\n".join(text_parts)
    return raw_text, None


# ── Text cleaning ─────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """
    Normalize extracted text for consistent downstream processing.

    Cleaning steps (in order):
        1. Remove non-printable / control characters except newlines and tabs
        2. Normalize Windows line endings to Unix (\r\n -> \n)
        3. Collapse runs of more than 2 consecutive newlines into exactly 2
        4. Collapse runs of spaces/tabs on a single line into one space
        5. Strip leading and trailing whitespace from the whole string

    We intentionally keep single and double newlines because they
    mark section boundaries that spaCy uses for section detection.

    Args:
        text: raw extracted string from PDF or DOCX

    Returns:
        Cleaned string ready for NLP processing
    """
    if not text:
        return ""

    # Step 1 — Remove non-printable control characters (keep \n and \t)
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)

    # Step 2 — Normalize Windows line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Step 3 — Collapse 3+ consecutive newlines into 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Step 4 — Collapse multiple spaces/tabs on the same line into one space
    text = re.sub(r"[ \t]+", " ", text)

    # Step 5 — Strip leading and trailing whitespace
    text = text.strip()

    return text


# ── Validation ────────────────────────────────────────────────────────────────

def _validate_text(text: str, ext: str) -> tuple:
    """
    Check that the extracted text has enough content to be analysed.

    The most common failure case is a scanned PDF — PyMuPDF returns
    empty text because there is no text layer, only images. We catch
    this here and return a clear, actionable error message.

    Args:
        text: cleaned extracted string
        ext:  file extension, used to give a specific error message

    Returns:
        (True, None) or (False, error_message)
    """
    if not text or len(text) < MIN_MEANINGFUL_TEXT_LENGTH:
        if ext == "pdf":
            return False, (
                "No text could be extracted from the PDF. "
                "It appears to be a scanned image. "
                "Please upload a text-based PDF or convert it to DOCX."
            )
        return False, (
            "No text could be extracted from the file. "
            "Please ensure the file contains readable text content."
        )

    return True, None


# ── Helper ────────────────────────────────────────────────────────────────────

def _get_extension(file_path: str) -> str:
    """
    Extract the lowercase extension from a file path.
    Returns empty string if no extension exists.

    Args:
        file_path: full file path string

    Returns:
        Lowercase extension without dot, e.g. 'pdf' or 'docx'
    """
    if "." not in file_path:
        return ""
    return file_path.rsplit(".", 1)[-1].lower()

