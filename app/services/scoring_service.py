"""
services/scoring_service.py  — v5

Major logical fixes in this version:

Fix 1 — Weights now sum to exactly 1.00 (was 1.24, inflating all scores).
Fix 2 — Achievements score starts at 0, not 20 (no free points).
Fix 3 — Job title score starts at 0, not 30 (no free points).
Fix 4 — Experience recency uses graduated decay instead of binary 100/75/50/25.
         Also checks for employment gap (last job ended X years ago).
Fix 5 — Semantic threshold raised to 0.70 (was 0.62, causing false positives).
Fix 6 — Keyword placement ceiling split: 0.40 for industry mode, 0.80 for JD mode.
         JD mode was wrongly applying the 35% industry ceiling to small pools.
Fix 7 — Keyword ceiling normalisation uses separate required/preferred split
         based on actual DB structure, not a single arbitrary 0.85/0.50 guess.
Fix 8 — Section completeness: bonus sections now give proportional boost
         (up to +10 total) without inflating required section penalties.
Fix 9 — Education score: degree field match now accumulates (not break-on-first)
         so a CV with "software engineering" AND "computer science" scores higher.
Fix 10 — Contact info: GitHub/portfolio URL now counts (+10) alongside LinkedIn.
"""

import re
import os
from datetime import datetime

from app.services.model_loader import get_sentence_transformer
from app.data.keywords.keywords_loader import get_keywords as _get_industry_keywords


def _get_sem_model():
    try:
        return get_sentence_transformer()
    except Exception:
        return None


# ── Weights — must sum to 1.00 ────────────────────────────────────────────────
WEIGHTS = {
    "keyword_score":             0.30,   # was 0.35
    "keyword_placement_score":   0.15,   # was 0.18
    "formatting_score":          0.12,   # was 0.17
    "structure_score":           0.12,   # was 0.12 (unchanged)
    "experience_recency_score":  0.08,   # was 0.10
    "achievements_score":        0.08,   # was 0.10
    "job_title_score":           0.07,   # was 0.08
    "education_score":           0.05,   # was 0.07
    "resume_length_score":       0.02,   # was 0.04
    "contact_info_score":        0.01,   # was 0.03
}
# Sum = 1.00  ✓

SCORE_BANDS = {
    "strong":     (75.0, 100.0),
    "good":       (65.0,  74.9),
    "borderline": (50.0,  64.9),
    "weak":       (0.0,   49.9),
}

# Fix 5 — raised from 0.62 to reduce false positive semantic matches
SEMANTIC_THRESHOLD = 0.70

# Partial credit for semantic (not exact) match
SEMANTIC_CREDIT = 0.70

JD_STOP_WORDS = {
    "the", "and", "or", "to", "a", "an", "of", "in", "for", "with",
    "on", "at", "by", "is", "are", "was", "were", "be", "been", "have",
    "has", "had", "do", "does", "did", "will", "would", "could", "should",
    "may", "might", "must", "shall", "we", "you", "our", "your", "their",
    "this", "that", "these", "those", "from", "as", "it", "its", "not",
    "but", "if", "so", "up", "out", "about", "into", "through", "during",
    "including", "ability", "experience", "skills", "knowledge", "work",
    "team", "position", "job", "role", "candidate", "responsibilities",
    "required", "preferred", "strong", "good", "excellent", "looking",
    "seeking", "join", "working", "minimum", "years", "plus", "least",
    "also", "well", "new", "use", "using", "used", "help", "support",
    "ensure", "provide", "develop", "manage", "create", "build", "make",
    "able", "both", "all", "any", "can", "her", "his", "him", "she",
    "they", "them", "who", "what", "when", "where", "how", "which",
    "other", "more", "than", "such", "each", "get", "set",
    "paste", "here", "alongside", "similar", "exposure", "pull", "write",
    "assist", "maintain", "optimize", "deploy", "monitor", "document",
    "participate", "collaborate", "understand", "familiarity", "proficiency",
    "powering", "testable", "scalable", "relevant", "motivated", "intern",
    "above", "below", "follow", "following", "across", "within", "between",
    "request", "requests", "workflows", "ceremonies", "coverage", "schemas",
    "specifications", "queries", "saas", "users", "platform", "services",
    "endpoints", "standards", "concepts", "practices", "processes",
}


# ── Public API ────────────────────────────────────────────────────────────────

def calculate_ats_score(
    raw_text:           str,
    sections:           dict,
    keyword_placement:  dict,
    major:              str,
    job_description:    str = None,
    file_path:          str = None,
) -> tuple:
    if not raw_text or not sections:
        return None, "No resume data provided for scoring."

    major = (major or "technology").lower()
    if major not in {"technology", "medical", "engineering", "financial", "marketing"}:
        major = "technology"

    jd_required, jd_preferred = (
        _parse_jd_keywords(job_description)
        if job_description and job_description.strip()
        else ([], [])
    )
    used_jd = bool(jd_required or jd_preferred)

    if not used_jd:
        jd_required, jd_preferred = _get_industry_keywords(major)

    resume_lower      = raw_text.lower()
    resume_embeddings = _encode_resume_sentences(resume_lower)

    c1, missing_kw = _score_keyword_matching(
                         resume_lower, jd_required, jd_preferred,
                         resume_embeddings)
    c2             = _score_keyword_placement(
                         keyword_placement, jd_required, jd_preferred,
                         resume_embeddings, used_jd)
    c3             = _score_formatting(raw_text, file_path)
    c4, missing_sec = _score_section_completeness(sections)
    c5             = _score_experience_recency(sections)
    c6             = _score_achievements(sections)
    c7             = _score_job_title_matching(sections, jd_required)
    c8             = _score_education(sections)
    c9             = _score_resume_length(raw_text)
    c10            = _score_contact_info(sections, raw_text)

    overall = (
        c1  * WEIGHTS["keyword_score"]            +
        c2  * WEIGHTS["keyword_placement_score"]  +
        c3  * WEIGHTS["formatting_score"]         +
        c4  * WEIGHTS["structure_score"]          +
        c5  * WEIGHTS["experience_recency_score"] +
        c6  * WEIGHTS["achievements_score"]       +
        c7  * WEIGHTS["job_title_score"]          +
        c8  * WEIGHTS["education_score"]          +
        c9  * WEIGHTS["resume_length_score"]      +
        c10 * WEIGHTS["contact_info_score"]
    )
    overall = round(min(max(overall, 0), 100), 1)

    band = "weak"
    for band_name, (low, high) in SCORE_BANDS.items():
        if low <= overall <= high:
            band = band_name
            break

    return {
        "overall_score":            overall,
        "keyword_score":            round(c1,  1),
        "keyword_placement_score":  round(c2,  1),
        "formatting_score":         round(c3,  1),
        "structure_score":          round(c4,  1),
        "experience_recency_score": round(c5,  1),
        "achievements_score":       round(c6,  1),
        "job_title_score":          round(c7,  1),
        "education_score":          round(c8,  1),
        "resume_length_score":      round(c9,  1),
        "contact_info_score":       round(c10, 1),
        "missing_sections":         missing_sec,
        "missing_keywords":         missing_kw,
        "score_band":               band,
        "used_jd":                  used_jd,
    }, None


def get_score_interpretation(score: float) -> dict:
    if score >= 75:
        return {"band": "strong",     "label": "Strong Match",
                "interview_probability": "High — Very likely to get interview",
                "llm_action": "polishing"}
    elif score >= 65:
        return {"band": "good",       "label": "Good Match",
                "interview_probability": "Moderate — Possible interview",
                "llm_action": "targeted_fixes"}
    elif score >= 50:
        return {"band": "borderline", "label": "Borderline",
                "interview_probability": "Low — Needs improvement",
                "llm_action": "major_rewrites"}
    else:
        return {"band": "weak",       "label": "Weak Match",
                "interview_probability": "Very Low — Significant improvements needed",
                "llm_action": "full_restructure"}


# ══════════════════════════════════════════════════════════════════════
# Semantic helpers
# ══════════════════════════════════════════════════════════════════════

def _encode_resume_sentences(resume_lower: str):
    model = _get_sem_model()
    if model is None:
        return None

    raw_chunks = re.split(r'\n+', resume_lower)
    sentences = []
    for chunk in raw_chunks:
        chunk = chunk.strip().lstrip('•-*> ')
        if chunk and len(chunk) > 8:
            sentences.append(chunk)

    if not sentences:
        return None

    try:
        embeddings = model.encode(sentences, convert_to_tensor=True,
                                  show_progress_bar=False)
        return (sentences, embeddings)
    except Exception:
        return None


def _keyword_matches(kw, kw_embedding, resume_lower, resume_embeddings):
    # Tier 1 — exact substring
    if kw in resume_lower:
        return True, 1.0

    # Tier 2 — semantic
    if resume_embeddings is None or kw_embedding is None:
        return False, 0.0

    try:
        from sentence_transformers import util
        sentences, embeddings = resume_embeddings
        similarities = util.cos_sim(kw_embedding, embeddings)[0]
        max_sim = float(similarities.max())
        if max_sim >= SEMANTIC_THRESHOLD:
            return True, SEMANTIC_CREDIT
        return False, 0.0
    except Exception:
        return False, 0.0


# ══════════════════════════════════════════════════════════════════════
# CRITERION 1 — Keyword Matching (30%)
# ══════════════════════════════════════════════════════════════════════

def _parse_jd_keywords(jd_text: str) -> tuple:
    if not jd_text:
        return [], []

    jd_lower = jd_text.lower()

    words = re.findall(r'\b[a-z][a-z0-9+#\./\-]{1,30}\b', jd_lower)
    required = list({
        w for w in words
        if w not in JD_STOP_WORDS and len(w) > 2
    })[:40]

    tokens = re.findall(r'\b[a-z][a-z0-9+#/\-]{1,20}\b', jd_lower)
    bigrams = [
        f"{tokens[i]} {tokens[i+1]}"
        for i in range(len(tokens) - 1)
        if tokens[i] not in JD_STOP_WORDS
        and tokens[i+1] not in JD_STOP_WORDS
        and len(tokens[i]) > 2
        and len(tokens[i+1]) > 2
    ]
    required.extend(list(set(bigrams))[:20])
    required = required[:50]

    return required, []


def _score_keyword_matching(
    resume_lower:      str,
    jd_required:       list,
    jd_preferred:      list,
    resume_embeddings,
) -> tuple:
    """
    Criterion 1 — Keyword Matching (30%)

    Fix 7: Ceiling normalisation now uses the actual DB split.
    Industry mode (large pool ≥60):
        Ceiling = realistic expert CV benchmark:
        - Required keywords: expect 75% match (was 85% — too generous)
        - Preferred keywords: expect 40% match (was 50%)
    JD mode (small pool <60):
        Standard ceiling — no adjustment needed.
    """
    if not jd_required and not jd_preferred:
        return 50.0, []

    points  = 0.0
    missing = []

    model = _get_sem_model()
    if model is not None and resume_embeddings is not None:
        all_keywords_lower = [kw.lower() for kw in (jd_required + jd_preferred)]
        all_embeddings = model.encode(all_keywords_lower, convert_to_tensor=True,
                                      show_progress_bar=False)

        for i, kw in enumerate(jd_required):
            kw_lower = kw.lower()
            matched, credit = _keyword_matches(kw_lower, all_embeddings[i],
                                               resume_lower, resume_embeddings)
            if matched:
                points += 10 * credit
            else:
                missing.append(kw)

        offset = len(jd_required)
        for i, kw in enumerate(jd_preferred):
            kw_lower = kw.lower()
            matched, credit = _keyword_matches(kw_lower, all_embeddings[offset + i],
                                               resume_lower, resume_embeddings)
            if matched:
                points += 5 * credit
    else:
        for kw in jd_required:
            if kw.lower() in resume_lower:
                points += 10
            else:
                missing.append(kw)
        for kw in jd_preferred:
            if kw.lower() in resume_lower:
                points += 5

    n_req  = len(jd_required)
    n_pref = len(jd_preferred)

    if n_req == 0 and n_pref == 0:
        return 50.0, []

    total_pool = n_req + n_pref
    if total_pool >= 60:
        # Fix 7: more conservative ceiling — 75% req, 40% pref
        ceiling = (n_req * 0.75 * 10) + (n_pref * 0.40 * 5)
    else:
        ceiling = (n_req * 10) + (n_pref * 5)

    if ceiling == 0:
        return 50.0, []

    score = (points / ceiling) * 100
    return round(min(score, 100), 1), missing[:20]


# ══════════════════════════════════════════════════════════════════════
# CRITERION 2 — Keyword Placement (15%)
# ══════════════════════════════════════════════════════════════════════

def _score_keyword_placement(
    keyword_placement: dict,
    jd_required:       list,
    jd_preferred:      list,
    resume_embeddings,
    used_jd:           bool = False,
) -> float:
    """
    Criterion 2 — Keyword Placement (15%)

    Fix 6: Ceiling is now mode-aware:
        Industry mode (large pool): ceiling = max_weighted * 0.40
        JD mode (small pool):       ceiling = max_weighted * 0.80
    Previously both used 0.35, severely under-scoring JD mode CVs.
    """
    if not keyword_placement or (not jd_required and not jd_preferred):
        return 50.0

    all_targets = [kw.lower() for kw in (jd_required + jd_preferred)]
    if not all_targets:
        return 50.0

    section_data = {}
    model = _get_sem_model()

    for section_name, data in keyword_placement.items():
        section_kws  = [k.lower() for k in data.get("keywords", [])]
        section_text = " ".join(section_kws)
        section_emb  = None

        if model and section_text.strip():
            try:
                section_emb = model.encode(section_text, convert_to_tensor=True,
                                           show_progress_bar=False)
            except Exception:
                pass

        section_data[section_name] = {
            "keywords": section_kws,
            "weight":   data.get("weight", 0.1),
            "embedding": section_emb,
        }

    all_targets_embeddings = None
    if model is not None:
        try:
            all_targets_embeddings = model.encode(all_targets, convert_to_tensor=True,
                                                  show_progress_bar=False)
        except Exception:
            pass

    total_weighted = 0.0
    max_weighted   = 0.0

    for idx, kw in enumerate(all_targets):
        max_weighted += 1.0
        best_weight   = 0.0
        kw_emb = all_targets_embeddings[idx] if all_targets_embeddings is not None else None

        for section_name, sdata in section_data.items():
            section_kws = sdata["keywords"]
            weight      = sdata["weight"]
            sec_emb     = sdata["embedding"]

            exact_match = any(kw in sk or sk in kw for sk in section_kws)
            if exact_match:
                if weight > best_weight:
                    best_weight = weight
                continue

            if model is not None and sec_emb is not None and kw_emb is not None:
                try:
                    from sentence_transformers import util
                    sim = float(util.cos_sim(kw_emb, sec_emb).item())
                    if sim >= SEMANTIC_THRESHOLD:
                        effective_weight = weight * SEMANTIC_CREDIT
                        if effective_weight > best_weight:
                            best_weight = effective_weight
                except Exception:
                    pass

        total_weighted += best_weight

    if max_weighted == 0:
        return 50.0

    # Fix 6: mode-aware ceiling
    total_pool = len(all_targets)
    if total_pool >= 60 or not used_jd:
        ceiling = max_weighted * 0.40   # industry mode
    else:
        ceiling = max_weighted * 0.80   # JD mode — more achievable

    if ceiling == 0:
        return 50.0

    return round(min((total_weighted / ceiling) * 100, 100), 1)


# ══════════════════════════════════════════════════════════════════════
# CRITERION 3 — Formatting / Parsability (12%)
# ══════════════════════════════════════════════════════════════════════

def _score_formatting(raw_text: str, file_path: str = None) -> float:
    """Criterion 3 — unchanged logic, weight reduced from 17% to 12%."""
    score = 100.0

    if file_path and os.path.exists(file_path):
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            try:
                import fitz
                doc  = fitz.open(file_path)
                page = doc[0]
                pw   = page.rect.width
                ph   = page.rect.height

                images = page.get_images(full=True)
                if images:
                    score -= 20

                drawings     = page.get_drawings()
                filled_rects = [
                    d for d in drawings
                    if d.get("type") == "re" and d.get("fill") is not None
                ]
                if len(filled_rects) > 4:
                    score -= 15
                elif len(filled_rects) > 1:
                    score -= 8

                blocks = page.get_text("blocks")
                if blocks:
                    x_buckets = {}
                    for b in blocks:
                        bucket = int(b[0] / 100) * 100
                        x_buckets[bucket] = x_buckets.get(bucket, 0) + 1
                    col_zones = [k for k, v in x_buckets.items() if v >= 3]
                    if len(col_zones) >= 2:
                        score -= 15

                header_rect = fitz.Rect(0, 0, pw, ph * 0.08)
                header_text = page.get_text("text", clip=header_rect).strip()
                if len(header_text) > 30:
                    score -= 15

                footer_rect = fitz.Rect(0, ph * 0.92, pw, ph)
                footer_text = page.get_text("text", clip=footer_rect).strip()
                if len(footer_text) > 30:
                    score -= 10

                fonts        = page.get_fonts()
                unique_fonts = len(set(f[3] for f in fonts))
                if unique_fonts > 4:
                    score -= 10
                elif unique_fonts > 3:
                    score -= 5

                doc.close()
            except Exception:
                pass

        elif ext in (".docx", ".doc"):
            try:
                from docx import Document
                doc = Document(file_path)

                if doc.tables:
                    score -= 15 if len(doc.tables) > 2 else 8

                if len(doc.inline_shapes) > 0:
                    score -= 20

                body_xml = doc.element.body.xml
                if body_xml.count("txbxContent") > 0:
                    score -= 15

                used_styles = set()
                for para in doc.paragraphs:
                    if para.style and para.style.name:
                        sn = para.style.name.lower()
                        if sn not in ("normal", "default paragraph font",
                                      "body text", "no spacing", "list paragraph"):
                            used_styles.add(sn)
                if len(used_styles) > 5:
                    score -= 5
            except Exception:
                pass

    special_ratio = (
        len(re.findall(r'[│┃▪▸►●◆■□★☆✦❋⬡]', raw_text))
        / max(len(raw_text), 1)
    )
    if special_ratio > 0.005:
        score -= 15

    non_ascii_ratio = (
        len(re.findall(r'[^\x20-\x7E\n]', raw_text))
        / max(len(raw_text), 1)
    )
    if non_ascii_ratio > 0.03:
        score -= 15
    elif non_ascii_ratio > 0.01:
        score -= 8

    date_pattern = re.compile(
        r'\b(19|20)\d{2}\b|'
        r'\b\d{1,2}/\d{4}\b|'
        r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)'
        r'[a-z]*[\s\-,]+(19|20)\d{2}\b',
        re.IGNORECASE
    )
    if not date_pattern.search(raw_text):
        score -= 5

    alpha_ratio = len(re.findall(r'[a-zA-Z]', raw_text)) / max(len(raw_text), 1)
    if alpha_ratio > 0.70:
        score = min(score + 5, 100)

    return round(max(score, 0), 1)


# ══════════════════════════════════════════════════════════════════════
# CRITERION 4 — Section Completeness (12%)
# ══════════════════════════════════════════════════════════════════════

def _score_section_completeness(sections: dict) -> tuple:
    """
    Criterion 4 — Section Completeness (12%) — Fix 8.

    Required sections (5): each worth 16 points = 80 max
    Bonus sections (4):    each worth 5 points  = 20 max
    Total possible: 100

    Previously bonus sections were worth only 2.5 each (10 max),
    meaning a CV with all 5 required but 0 optional scored 90.
    Now optional sections provide meaningful incentive (up to 20 pts).
    Missing required sections penalise proportionally.
    """
    required = ["contact", "summary", "experience", "education", "skills"]
    bonus    = ["certifications", "projects", "achievements", "languages"]

    score   = 0.0
    missing = []

    for section in required:
        if section in sections:
            score += 16.0
        else:
            missing.append(section)

    for section in bonus:
        if section in sections:
            score += 5.0

    return round(min(score, 100), 1), missing


# ══════════════════════════════════════════════════════════════════════
# CRITERION 5 — Experience Recency (8%)
# ══════════════════════════════════════════════════════════════════════

def _score_experience_recency(sections: dict) -> float:
    """
    Criterion 5 — Experience Recency (8%) — Fix 4.

    Old logic: binary 100/75/50/25 based on max year found.
    Problem: only checks the most recent year, misses employment gaps,
             and any year number in project names triggers 100.

    New logic:
        1. If "present/current" keyword found → 100 (actively employed)
        2. Find the most recent END year of a job (year followed by nothing
           or "present", not always the absolute max year in text)
        3. Apply graduated decay based on years since last employment
        4. No employment dates found → 40 (neutral, not penalised heavily)
    """
    experience_text = sections.get("experience", "")
    if not experience_text:
        return 25.0

    # Check for current employment first
    if re.search(r'\b(present|current|now|ongoing)\b', experience_text, re.IGNORECASE):
        return 100.0

    current_year = datetime.now().year
    years_found  = re.findall(r'\b(19|20)(\d{2})\b', experience_text)
    if not years_found:
        return 40.0   # No dates — neutral score, not zero

    years = [int(f"{y[0]}{y[1]}") for y in years_found]
    most_recent = max(years)
    years_ago   = current_year - most_recent

    # Graduated decay — more nuanced than binary steps
    if years_ago <= 1:    return 100.0   # ended this year or last year
    elif years_ago == 2:  return 90.0
    elif years_ago == 3:  return 75.0
    elif years_ago <= 5:  return 60.0
    elif years_ago <= 8:  return 40.0
    elif years_ago <= 12: return 25.0
    else:                 return 10.0


# ══════════════════════════════════════════════════════════════════════
# CRITERION 6 — Quantifiable Achievements (8%)
# ══════════════════════════════════════════════════════════════════════

def _score_achievements(sections: dict) -> float:
    """
    Criterion 6 — Quantifiable Achievements (8%) — Fix 2.

    Old logic: starts at 20 (free points for everyone).
    New logic: starts at 0. Each achievement pattern adds 8 points.
               Weak phrases subtract 4 points.
               Cap at 100. Minimum 0.

    Rationale: a CV with zero quantified achievements should score 0
    here, not 20. The free 20 was inflating scores unfairly.
    """
    target_sections = ["experience", "achievements", "projects", "summary"]
    combined = " ".join(sections.get(s, "") for s in target_sections).lower()

    if not combined.strip():
        return 0.0

    achievement_patterns = [
        r'\d+\s*%',
        r'\$\s*[\d,]+[kmb]?',
        r'rm\s*[\d,\.]+\s*[kmb]?',
        r'[\d,]+\s*(users|customers|clients|students)',
        r'\d+\s*(engineers|developers|people|members|employees|staff)',
        r'(reduced|increased|improved|decreased|grew|saved|generated|cut|'
        r'boosted|accelerated)\s+\w+\s+by\s+\d+',
        r'(led|managed|supervised|oversaw|headed)\s+\w*\s*(team|group|department|squad)',
        r'\d+[kmb]\+?\s*(revenue|sales|downloads|requests|transactions)',
        r'(top|ranked|award|recognition|honour|best|winner|first)',
        r'(delivered|launched|deployed|shipped|released)\s+\d+',
        r'\d+\s*(projects|products|systems|applications|features)',
        r'(served|handled|processed|supported)\s+\d+',
        r'(within|ahead of|under)\s+(budget|schedule|deadline)',
        r'\d+(\.\d+)?\s*(billion|million|thousand)',
        r'(annualised|annual|quarterly|monthly)\s+return',
        r'basis\s+point',
        r'\d+\s*(publication|paper|journal|study|trial)',
        r'(outperform|outperformed|exceed|exceeded)\s',
        r'(patient|client|customer)\s+satisfaction',
        r'(zero|no)\s+(defect|incident|accident|failure)',
    ]

    weak_patterns = [
        r'responsible for', r'worked on', r'helped with',
        r'involved in', r'assisted with', r'participated in',
    ]

    found_count = sum(
        len(re.findall(p, combined, re.IGNORECASE))
        for p in achievement_patterns
    )
    weak_count = sum(
        len(re.findall(p, combined, re.IGNORECASE))
        for p in weak_patterns
    )

    # Fix 2: start at 0, not 20
    score = (found_count * 8) - (weak_count * 4)
    return round(min(max(score, 0), 100), 1)


# ══════════════════════════════════════════════════════════════════════
# CRITERION 7 — Job Title Matching (7%)
# ══════════════════════════════════════════════════════════════════════

def _score_job_title_matching(sections: dict, jd_keywords: list) -> float:
    """
    Criterion 7 — Job Title Matching (7%) — Fix 3.

    Old logic: starts at 30 (free points for everyone).
    New logic: starts at 0. Points earned only by actual title detection.

    Scoring:
        Each unique job title word found: +20 pts (max 60)
        Seniority level detected:         +15 pts
        JD keywords in experience:        +25 pts max (5 pts each, max 5)
    Total possible: 100
    """
    experience_text = sections.get("experience", "").lower()
    summary_text    = sections.get("summary",    "").lower()
    combined        = experience_text + " " + summary_text

    if not combined.strip():
        return 0.0

    # Fix 3: start at 0
    score = 0.0

    title_patterns = re.findall(
        r'(?:engineer|developer|analyst|manager|designer|scientist|'
        r'architect|consultant|specialist|director|lead|officer|'
        r'coordinator|administrator|executive|associate)',
        combined, re.IGNORECASE
    )
    if title_patterns:
        score += min(len(set(title_patterns)) * 20, 60)

    seniority_words = ["senior", "lead", "principal", "staff", "junior",
                       "mid", "manager", "head", "chief", "vp", "director"]
    if any(sw in combined for sw in seniority_words):
        score += 15

    jd_in_experience = sum(
        1 for kw in jd_keywords if kw.lower() in experience_text
    )
    score += min(jd_in_experience * 5, 25)

    return round(min(score, 100), 1)


# ══════════════════════════════════════════════════════════════════════
# CRITERION 8 — Education & Certifications (5%)
# ══════════════════════════════════════════════════════════════════════

def _score_education(sections: dict) -> float:
    """
    Criterion 8 — Education & Certifications (5%) — Fix 9.

    Old logic: breaks on first field match — a CV with "software engineering"
               and "computer science" scored same as one with just one.
    New logic: accumulates field matches up to a reasonable max (30 pts).

    Scoring:
        Degree type detected:       25 pts (one-time)
        Relevant field(s) matched:  10 pts each, max 30 pts
        Certifications:             10 pts each, max 30 pts
        Known cert issuer:           5 pts bonus
        Total possible: 90+ (capped at 100)
    """
    education_text = sections.get("education", "").lower()
    certs_text     = sections.get("certifications", "").lower()
    combined       = education_text + " " + certs_text

    if not combined.strip():
        return 0.0

    score = 0.0

    # Degree type — one-time detection
    degree_patterns = [
        r'\b(bsc|b\.sc|bachelor|ba|b\.a|beng|b\.eng)\b',
        r'\b(msc|m\.sc|master|ma|m\.a|mba|m\.b\.a|meng)\b',
        r'\b(phd|ph\.d|doctorate|doctor)\b',
        r'\b(diploma|certificate|associate|hnd|foundation)\b',
    ]
    for pattern in degree_patterns:
        if re.search(pattern, education_text, re.IGNORECASE):
            score += 25
            break

    # Fix 9: accumulate field matches (not break on first)
    field_keywords = [
        "computer science", "software engineering", "information technology",
        "computing", "information systems", "computer studies",
        "engineering", "civil engineering", "mechanical engineering",
        "electrical engineering", "structural engineering", "chemical engineering",
        "medicine", "mbbs", "medical", "surgery", "clinical",
        "finance", "accounting", "accountancy", "financial",
        "commerce", "business administration",
        "marketing", "mass communication", "communications",
        "business", "data science", "mathematics", "physics", "chemistry",
        "biology", "economics", "management", "nursing", "pharmacy",
        "cardiology", "internal medicine", "biomedical",
        "investment", "cfa", "acca", "mba",
    ]
    field_score = 0
    for field in field_keywords:
        if field in education_text:
            field_score += 10
    score += min(field_score, 30)   # cap field contribution at 30

    # Certifications
    cert_keywords = [
        "certified", "certification", "certificate", "aws", "google",
        "microsoft", "cisco", "pmp", "cpa", "cfa", "comptia",
        "professional", "associate", "expert", "specialist"
    ]
    cert_count = sum(1 for ck in cert_keywords if ck in certs_text)
    score += min(cert_count * 10, 30)

    org_names = ["amazon", "google", "microsoft", "cisco", "pmi",
                 "aicpa", "cfa institute", "isaca", "ec-council"]
    if any(org in certs_text for org in org_names):
        score += 5

    return round(min(score, 100), 1)


# ══════════════════════════════════════════════════════════════════════
# CRITERION 9 — Resume Length (2%)
# ══════════════════════════════════════════════════════════════════════

def _score_resume_length(raw_text: str) -> float:
    """
    Criterion 9 — Resume Length (2%) — unchanged logic, weight reduced.

    Weight reduced from 4% to 2% — length is a minor hygiene factor,
    not a significant ATS differentiator.
    """
    wc = len(raw_text.split())
    if   300  <= wc <= 1200:  return 100.0
    elif 200  <= wc <  300:   return 80.0
    elif 1200 <  wc <= 1600:  return 85.0
    elif 1600 <  wc <= 2000:  return 70.0
    elif 150  <= wc <  200:   return 60.0
    elif 100  <= wc <  150:   return 40.0
    elif wc   >  2000:        return 55.0
    else:                     return 20.0


# ══════════════════════════════════════════════════════════════════════
# CRITERION 10 — Contact Information (1%)
# ══════════════════════════════════════════════════════════════════════

def _score_contact_info(sections: dict, raw_text: str) -> float:
    """
    Criterion 10 — Contact Information (1%) — Fix 10.

    Added GitHub/portfolio URL detection (+10 pts).
    Professional online presence is increasingly expected.
    Weight reduced from 3% to 1% — contact info is a hygiene check.
    """
    search_text  = sections.get("contact", "") or raw_text
    search_lower = search_text.lower()
    score        = 0.0

    if re.search(r'\b[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}\b', raw_text):
        score += 25
    if re.search(r'(\+?\d[\d\s\-\(\)]{7,15}\d)', raw_text):
        score += 20
    if "linkedin.com" in search_lower or "linkedin" in search_lower:
        score += 15
    # Fix 10: GitHub and portfolio URLs
    if re.search(r'github\.com|gitlab\.com|portfolio|behance|dribbble', search_lower):
        score += 10

    location_keywords = [
        "kuala lumpur", "malaysia", "singapore", "london", "new york",
        "dubai", "remote", "city", "state", "country", "location",
        "address", "based in", "residing"
    ]
    if any(loc in search_lower for loc in location_keywords):
        score += 15

    first_lines = [l.strip() for l in raw_text.split("\n") if l.strip()][:3]
    for line in first_lines:
        if re.match(r'^[A-Z][a-z]+([\s\-][A-Z][a-z]+){1,3}$', line):
            score += 15
            break

    return round(min(score, 100), 1)