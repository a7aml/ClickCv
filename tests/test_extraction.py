"""
tests/test_extraction.py

Tests for services/extraction_service.py

Run from project root:
    python tests/test_extraction.py
"""

import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import tempfile
import fitz
from docx import Document
from app.services.extraction_service import extract_text_from_file


# ── Helpers ───────────────────────────────────────────────────────────────────

def make_temp_pdf(text: str, password: str = None) -> str:
    """
    Create a real minimal PDF with the given text.
    Uses delete=False + explicit close to avoid Windows file lock.
    PyMuPDF cannot write to a file that Python still holds open.
    """
    # Step 1 — create the temp file and immediately close it
    # so Windows releases the file handle before PyMuPDF writes to it
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp_path = tmp.name
    tmp.close()   # Release handle — critical on Windows

    # Step 2 — now PyMuPDF can write freely
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((50, 50), text, fontsize=11)

    if password:
        pdf.save(tmp_path, encryption=fitz.PDF_ENCRYPT_AES_256,
                 user_pw=password, owner_pw=password)
    else:
        pdf.save(tmp_path)

    pdf.close()
    return tmp_path


def make_temp_docx(paragraphs: list, table_content: list = None) -> str:
    """
    Create a real minimal DOCX with the given paragraphs.
    Uses delete=False + explicit close for Windows compatibility.
    """
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp_path = tmp.name
    tmp.close()   # Release handle before python-docx writes

    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)

    if table_content:
        table = doc.add_table(rows=len(table_content), cols=1)
        for i, content in enumerate(table_content):
            table.rows[i].cells[0].text = content

    doc.save(tmp_path)
    return tmp_path


def make_empty_pdf() -> str:
    """Create a real PDF with no text — simulates a scanned image PDF."""
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp_path = tmp.name
    tmp.close()

    pdf = fitz.open()
    pdf.new_page()   # Empty page — no text inserted
    pdf.save(tmp_path)
    pdf.close()
    return tmp_path


# ── Test runner ───────────────────────────────────────────────────────────────

def run_tests():
    print("Running extraction_service tests...\n")
    passed = 0
    failed = 0
    temp_files = []

    def check(label, file_path, expected_success,
              expected_text_contains=None, expected_error_contains=None):
        nonlocal passed, failed

        text, error = extract_text_from_file(file_path)

        if expected_success and text is None:
            print(f"  FAIL  {label}")
            print(f"        expected text, got error: {error}")
            failed += 1
            return

        if not expected_success and error is None:
            print(f"  FAIL  {label}")
            print(f"        expected error, got text: {text[:60]}...")
            failed += 1
            return

        if expected_text_contains and expected_text_contains.lower() not in (text or "").lower():
            print(f"  FAIL  {label}")
            print(f"        expected text to contain '{expected_text_contains}'")
            print(f"        got: {(text or '')[:100]}")
            failed += 1
            return

        if expected_error_contains and expected_error_contains.lower() not in (error or "").lower():
            print(f"  FAIL  {label}")
            print(f"        expected error containing '{expected_error_contains}'")
            print(f"        got: {error}")
            failed += 1
            return

        print(f"  PASS  {label}")
        passed += 1

    # ── File path checks ──────────────────────────────────────────────────────
    check(
        "Non-existent file path",
        "/fake/path/resume.pdf",
        expected_success=False,
        expected_error_contains="not found"
    )

    check(
        "Unsupported file type .txt",
        "/fake/path/resume.txt",
        expected_success=False,
        expected_error_contains="Unsupported"
    )

    check(
        "Unsupported file type .exe",
        "/fake/path/resume.exe",
        expected_success=False,
        expected_error_contains="Unsupported"
    )

    # ── PDF extraction ────────────────────────────────────────────────────────
    cv_text = (
        "John Smith\n"
        "john@email.com | +1234567890\n\n"
        "EXPERIENCE\n"
        "Software Engineer at Google 2020-2023\n"
        "Developed scalable microservices using Python and Flask.\n\n"
        "EDUCATION\n"
        "BSc Computer Science MIT 2020\n\n"
        "SKILLS\n"
        "Python Flask PostgreSQL Docker AWS"
    )

    pdf_path = make_temp_pdf(cv_text)
    temp_files.append(pdf_path)

    check(
        "Valid PDF — extracts experience",
        pdf_path,
        expected_success=True,
        expected_text_contains="Software Engineer"
    )

    check(
        "Valid PDF — extracts skills",
        pdf_path,
        expected_success=True,
        expected_text_contains="Python"
    )

    check(
        "Valid PDF — extracts education",
        pdf_path,
        expected_success=True,
        expected_text_contains="Computer Science"
    )

    # ── Empty / scanned PDF ───────────────────────────────────────────────────
    empty_pdf_path = make_empty_pdf()
    temp_files.append(empty_pdf_path)

    check(
        "Empty PDF (simulates scanned image)",
        empty_pdf_path,
        expected_success=False,
        expected_error_contains="scanned"
    )

    # ── Password protected PDF ────────────────────────────────────────────────
    protected_pdf_path = make_temp_pdf(cv_text, password="secret123")
    temp_files.append(protected_pdf_path)

    check(
        "Password-protected PDF",
        protected_pdf_path,
        expected_success=False,
        expected_error_contains="password"
    )

    # ── DOCX extraction ───────────────────────────────────────────────────────
    paragraphs = [
        "Jane Doe",
        "jane@email.com | +9876543210",
        "EXPERIENCE",
        "Data Scientist at Amazon 2021-2024",
        "Built ML pipelines using TensorFlow and PyTorch.",
        "EDUCATION",
        "MSc Data Science Stanford University 2021",
        "SKILLS",
        "Python TensorFlow PyTorch SQL Tableau"
    ]

    docx_path = make_temp_docx(paragraphs)
    temp_files.append(docx_path)

    check(
        "Valid DOCX — extracts experience",
        docx_path,
        expected_success=True,
        expected_text_contains="Data Scientist"
    )

    check(
        "Valid DOCX — extracts skills",
        docx_path,
        expected_success=True,
        expected_text_contains="TensorFlow"
    )

    check(
        "Valid DOCX — extracts education",
        docx_path,
        expected_success=True,
        expected_text_contains="Stanford"
    )

    # ── DOCX with table content ───────────────────────────────────────────────
    table_paragraphs = ["Alice Brown", "alice@email.com"]
    table_cells = [
        "Project Manager",
        "PMP Certified",
        "Agile Scrum JIRA Confluence"
    ]

    docx_table_path = make_temp_docx(table_paragraphs, table_cells)
    temp_files.append(docx_table_path)

    check(
        "DOCX with table — extracts table rows",
        docx_table_path,
        expected_success=True,
        expected_text_contains="Project Manager"
    )

    check(
        "DOCX with table — extracts table cells",
        docx_table_path,
        expected_success=True,
        expected_text_contains="Agile"
    )

    # ── Text cleaning check ───────────────────────────────────────────────────
    messy_text = (
        "John   Smith\n"
        "Python    Flask    PostgreSQL\n"
        "Built   scalable   APIs   for   enterprise   clients\n"
        "BSc Computer Science 2020"
    )

    messy_pdf_path = make_temp_pdf(messy_text)
    temp_files.append(messy_pdf_path)

    text, error = extract_text_from_file(messy_pdf_path)
    if text:
        if "  " in text:
            print("  FAIL  Text cleaning — double spaces not collapsed")
            failed += 1
        else:
            print("  PASS  Text cleaning — double spaces collapsed correctly")
            passed += 1
    else:
        print(f"  FAIL  Text cleaning — extraction failed: {error}")
        failed += 1

    # ── Cleanup ───────────────────────────────────────────────────────────────
    for path in temp_files:
        try:
            os.unlink(path)
        except Exception:
            pass

    # ── Summary ───────────────────────────────────────────────────────────────
    print(f"\n{passed} passed — {failed} failed")


if __name__ == "__main__":
    run_tests()